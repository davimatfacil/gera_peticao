import streamlit as st
import pdfplumber
from docx import Document
import re
import io
import base64
from docx.shared import Pt

class ProcessadorHolerite:
    def __init__(self):
        self.dados_extraidos = {}
    
    def extrair_dados_pdf(self, caminho_pdf):
        """
        Extrai dados relevantes do holerite em PDF
        """
        try:
            with pdfplumber.open(caminho_pdf) as pdf:
                texto = ""
                for pagina in pdf.pages:
                    texto += pagina.extract_text() + "\n"
                
                # Melhorar a extração do nome para não incluir a matrícula
                nome_match = re.search(r"Nome:\s*(.*?)(?=\s*Matrícula:)", texto, re.DOTALL)
                nome = nome_match.group(1).strip() if nome_match else ""
                
                self.dados_extraidos = {
                    'nome_completo': nome,
                    'matricula': self._extrair_campo(texto, r"Matrícula:\s*(\d+)"),
                    'cargo': self._extrair_campo(texto, r"Cargo:\s*(.*?)\n"),
                    'mes_ano': self._extrair_campo(texto, r"Mês/Ano:\s*(.*?)\n"),
                    'salario_base': self._extrair_valor(texto, r"Salário Base.*?R\$\s*([\d\.,]+)"),
                    'valor_inss': self._extrair_valor(texto, r"INSS.*?R\$\s*([\d\.,]+)"),
                    'valor_irrf': self._extrair_valor(texto, r"IRRF.*?R\$\s*([\d\.,]+)"),
                    'salario_liquido': self._extrair_valor(texto, r"Valor Líquido:\s*R\$\s*([\d\.,]+)")
                }
                
                return self.dados_extraidos
                
        except Exception as e:
            st.error(f"Erro ao processar PDF: {str(e)}")
            return None
    
    def _extrair_campo(self, texto, padrao):
        match = re.search(padrao, texto)
        return match.group(1).strip() if match else ""
    
    def _extrair_valor(self, texto, padrao):
        match = re.search(padrao, texto)
        if match:
            valor_str = match.group(1).replace('.', '').replace(',', '.')
            return float(valor_str)
        return 0.0
    
    def _formatar_valor_monetario(self, valor):
        """
        Formata valor para o padrão brasileiro de moeda
        """
        if isinstance(valor, (int, float)):
            return f"R$ {valor:,.2f}".replace(',', '_').replace('.', ',').replace('_', '.')
        return valor
    
    def preencher_peticao(self, template_file, dados_adicionais=None):
        try:
            doc = Document(template_file)
            
            # Combinar dados extraídos com dados adicionais
            dados_completos = {
                **self.dados_extraidos,
                **(dados_adicionais or {})
            }
            
            # Formatar valores monetários
            dados_completos['salario_base'] = self._formatar_valor_monetario(dados_completos['salario_base'])
            dados_completos['valor_inss'] = self._formatar_valor_monetario(dados_completos['valor_inss'])
            dados_completos['valor_irrf'] = self._formatar_valor_monetario(dados_completos['valor_irrf'])
            dados_completos['salario_liquido'] = self._formatar_valor_monetario(dados_completos['salario_liquido'])
            
            # Lista para armazenar os parágrafos
            paragrafos = []
            for paragrafo in doc.paragraphs:
                texto_original = paragrafo.text
                texto_novo = texto_original
                
                # Lista de substituições para garantir a ordem correta
                substituicoes = [
                    ('[NOME_COMPLETO]', dados_completos.get('nome_completo', '')),
                    ('[SALARIO_BASE]', dados_completos.get('salario_base', '')),
                    ('[VALOR_INSS]', dados_completos.get('valor_inss', '')),
                    ('[VALOR_IRRF]', dados_completos.get('valor_irrf', '')),
                    ('[SALARIO_LIQUIDO]', dados_completos.get('salario_liquido', ''))
                ]
                
                # Aplicar todas as substituições
                for placeholder, valor in substituicoes:
                    if placeholder in texto_novo:
                        texto_novo = texto_novo.replace(placeholder, str(valor))
                
                # Aplicar outras substituições dos dados adicionais
                for chave, valor in dados_completos.items():
                    placeholder = f"[{chave.upper()}]"
                    if placeholder in texto_novo:
                        texto_novo = texto_novo.replace(placeholder, str(valor))
                
                paragrafos.append(texto_novo)
            
            return paragrafos
            
        except Exception as e:
            st.error(f"Erro ao preencher petição: {str(e)}")
            return None
            
    def gerar_documento_final(self, paragrafos_editados):
        """
        Gera documento final com os parágrafos editados
        """
        try:
            doc = Document()
            for texto in paragrafos_editados:
                if texto.strip():  # Ignorar parágrafos vazios
                    p = doc.add_paragraph(texto)
            
            # Salvar em memória
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            return doc_io
        except Exception as e:
            st.error(f"Erro ao gerar documento final: {str(e)}")
            return None

def get_download_link(buffer, filename):
    """Gera link para download do arquivo"""
    b64 = base64.b64encode(buffer.getvalue()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" class="download-button">Baixar Versão Final da Petição</a>'

def main():
    st.set_page_config(page_title="Gerador de Petições", page_icon="⚖️", layout="wide")
    
    # Inicializar estado da sessão
    if 'current_step' not in st.session_state:
        st.session_state.current_step = 1
    if 'paragrafos_gerados' not in st.session_state:
        st.session_state.paragrafos_gerados = None
    if 'dados_extraidos' not in st.session_state:
        st.session_state.dados_extraidos = None
    if 'paragrafos_editados' not in st.session_state:
        st.session_state.paragrafos_editados = None

    st.title("🔖 Gerador Automático de Petições")
    
    # Barra de progresso
    progresso = st.progress(0)
    st.write("---")

    # Passo 1: Upload e Extração
    if st.session_state.current_step == 1:
        progresso.progress(33)
        st.subheader("📄 Passo 1: Upload e Extração de Dados")
        
        col1, col2 = st.columns(2)
        
        with col1:
            uploaded_holerite = st.file_uploader("Upload do Holerite (PDF)", type="pdf")
            uploaded_template = st.file_uploader("Upload do Template da Petição (DOCX)", type="docx")
            
            with st.expander("➕ Dados Adicionais"):
                dados_adicionais = {
                    "nacionalidade": st.text_input("Nacionalidade", value="brasileiro"),
                    "estado_civil": st.text_input("Estado Civil", value="casado"),
                    "rg": st.text_input("RG", value="12.345.678-9"),
                    "cpf": st.text_input("CPF", value="123.456.789-00"),
                    "endereco": st.text_input("Endereço", value="Rua Exemplo, 123 - São Paulo/SP"),
                    "comarca": st.text_input("Comarca", value="SÃO PAULO"),
                    "nome_empresa": st.text_input("Nome da Empresa", value="Empresa ABC Ltda."),
                    "cnpj": st.text_input("CNPJ", value="12.345.678/0001-90"),
                    "endereco_empresa": st.text_input("Endereço da Empresa", value="Avenida Empresarial, 456 - São Paulo/SP"),
                    "data_admissao": st.text_input("Data de Admissão", value="01/01/2020")
                }

        with col2:
            if uploaded_holerite and uploaded_template:
                processador = ProcessadorHolerite()
                dados = processador.extrair_dados_pdf(uploaded_holerite)
                
                if dados:
                    st.success("✅ Dados extraídos com sucesso!")
                    st.session_state.dados_extraidos = dados
                    
                    # Mostrar dados extraídos
                    dados_formatados = []
                    for chave, valor in dados.items():
                        if isinstance(valor, (int, float)):
                            valor_formatado = processador._formatar_valor_monetario(valor)
                        else:
                            valor_formatado = valor
                        dados_formatados.append({"Campo": chave, "Valor": valor_formatado})
                    
                    st.table(dados_formatados)
                    
                    # Botão para gerar documento preliminar
                    if st.button("➡️ Gerar Documento Preliminar"):
                        paragrafos = processador.preencher_peticao(uploaded_template, dados_adicionais)
                        if paragrafos:
                            st.session_state.paragrafos_gerados = paragrafos
                            st.session_state.current_step = 2
                            st.rerun()

    # Passo 2: Revisão e Edição
    elif st.session_state.current_step == 2:
        progresso.progress(66)
        st.subheader("📝 Passo 2: Revisão e Edição do Documento")
        
        if st.session_state.paragrafos_gerados:
            st.write("### Revise e edite o documento:")
            
            # Criar lista para armazenar parágrafos editados
            paragrafos_editados = []
            
            # Interface de edição por parágrafos
            for i, texto in enumerate(st.session_state.paragrafos_gerados):
                if texto.strip():  # Ignorar parágrafos vazios
                    texto_editado = st.text_area(f"Parágrafo {i+1}", value=texto, height=100)
                    paragrafos_editados.append(texto_editado)
            
            col1, col2 = st.columns(2)
            with col1:
                if st.button("⬅️ Voltar"):
                    st.session_state.current_step = 1
                    st.rerun()
            with col2:
                if st.button("➡️ Finalizar Edição"):
                    # Salvar parágrafos editados
                    st.session_state.paragrafos_editados = paragrafos_editados
                    st.session_state.current_step = 3
                    st.rerun()

    # Passo 3: Finalização
    elif st.session_state.current_step == 3:
        progresso.progress(100)
        st.subheader("✅ Passo 3: Documento Final")
        
        if st.session_state.paragrafos_editados:
            processador = ProcessadorHolerite()
            doc_final = processador.gerar_documento_final(st.session_state.paragrafos_editados)
            
            if doc_final:
                st.success("Documento revisado e finalizado com sucesso!")
                
                # Opções finais
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("⬅️ Voltar para Edição"):
                        st.session_state.current_step = 2
                        st.rerun()
                
                with col2:
                    st.markdown(get_download_link(doc_final, 
                                                "peticao_final.docx"), 
                              unsafe_allow_html=True)

                # Adicionar CSS personalizado para o botão de download
                st.markdown("""
                    <style>
                    .download-button {
                        background-color: #4CAF50;
                        border: none;
                        color: white;
                        padding: 12px 24px;
                        text-align: center;
                        text-decoration: none;
                        display: inline-block;
                        font-size: 16px;
                        margin: 4px 2px;
                        cursor: pointer;
                        border-radius: 4px;
                    }
                    </style>
                    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()